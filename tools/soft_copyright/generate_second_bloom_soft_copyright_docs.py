from __future__ import annotations

import math
import re
import shutil
import subprocess
import textwrap
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


SOFTWARE_NAME = "Second Bloom旧衣新生系统"
VERSION = "V1.0"
COPYRIGHT_OWNER = "张子妍"
OWNER_CERTIFICATE = "待补充营业执照或身份证号"
COMPANY_FOUNDING_DATE = "2025年12月01日"
SOFTWARE_COMPLETION_DATE = "2026年03月05日"
SOFTWARE_PUBLICATION_DATE = "2026年04月16日"
FIRST_PUBLICATION_LOCATION = "中国"
CLIENT_MIN_OS = "Android 7.0 及以上"
BACKEND_RUNTIME = "Linux 云端服务环境"

SCREENSHOT_SPECS = [
    ("01_inspiration.png", "灵感空间首页"),
    ("02_upload_demo.png", "上传旧衣与示例图选择"),
    ("03_low_confidence.png", "低置信度提示"),
    ("04_analysis_summary.png", "识别结果确认与编辑"),
    ("05_plan_generation.png", "方案生成页"),
    ("06_preview_editor.png", "真图编辑页"),
    ("07_preview_result.png", "最终效果图页"),
    ("08_wardrobe.png", "数字衣橱"),
    ("09_planet.png", "可持续星球"),
    ("10_profile.png", "我的主页"),
    ("11_auth.png", "登录与账号页面"),
]

ANDROID_FILE_ORDER = [
    "app/src/main/java/com/scf/secondbloom/MainActivity.kt",
    "app/src/main/java/com/scf/secondbloom/ui/MainScreen.kt",
    "app/src/main/java/com/scf/secondbloom/navigation/Screen.kt",
    "app/src/main/java/com/scf/secondbloom/presentation/remodel/RemodelViewModel.kt",
    "app/src/main/java/com/scf/secondbloom/domain/model/AppLanguage.kt",
    "app/src/main/java/com/scf/secondbloom/domain/model/RemodelModels.kt",
    "app/src/main/java/com/scf/secondbloom/domain/model/RemodelDerivedModels.kt",
    "app/src/main/java/com/scf/secondbloom/ui/screens/HomeScreen.kt",
    "app/src/main/java/com/scf/secondbloom/ui/screens/WorkbenchScreen.kt",
    "app/src/main/java/com/scf/secondbloom/ui/screens/PreviewEditorScreen.kt",
    "app/src/main/java/com/scf/secondbloom/ui/screens/PreviewResultScreen.kt",
    "app/src/main/java/com/scf/secondbloom/ui/screens/InspirationScreen.kt",
    "app/src/main/java/com/scf/secondbloom/ui/screens/InspirationDetailScreen.kt",
    "app/src/main/java/com/scf/secondbloom/ui/screens/WardrobeScreen.kt",
    "app/src/main/java/com/scf/secondbloom/ui/screens/PlanetScreen.kt",
    "app/src/main/java/com/scf/secondbloom/ui/screens/ProfileScreen.kt",
    "app/src/main/java/com/scf/secondbloom/ui/screens/AuthScreen.kt",
    "app/src/main/java/com/scf/secondbloom/ui/screens/AccountScreen.kt",
    "app/src/main/java/com/scf/secondbloom/data/remote/RealRemodelApi.kt",
    "app/src/main/java/com/scf/secondbloom/data/remote/dto/RemodelDtos.kt",
    "app/src/main/java/com/scf/secondbloom/data/local/FileRemodelHistoryLocalDataSource.kt",
    "app/src/main/java/com/scf/secondbloom/data/local/DefaultRemodelHistoryRepository.kt",
    "app/src/main/java/com/scf/secondbloom/data/historysync/HistorySyncModels.kt",
    "app/src/main/java/com/scf/secondbloom/data/historysync/HistorySyncHttpClient.kt",
    "app/src/main/java/com/scf/secondbloom/auth/SecondBloomClerkConfig.kt",
    "app/src/main/java/com/scf/secondbloom/auth/ClerkHistoryAuthTokenProvider.kt",
]

BACKEND_FILE_ORDER = [
    "app/main.py",
    "app/api/routes/health.py",
    "app/api/routes/analyze.py",
    "app/api/routes/plans.py",
    "app/api/routes/previews.py",
    "app/api/routes/me.py",
    "app/api/services.py",
    "app/services/analyze_service.py",
    "app/services/plan_service.py",
    "app/services/preview_job_service.py",
    "app/services/preview_render_service.py",
    "app/providers/base.py",
    "app/providers/qwen_chat.py",
    "app/providers/qwen_image_edit.py",
    "app/providers/qwen_visual_qa.py",
    "app/db/storage.py",
    "app/db/repository.py",
    "app/db/models/analysis.py",
    "app/db/models/plan.py",
    "app/db/models/preview.py",
    "app/db/models/user.py",
]


@dataclass(frozen=True)
class RepoStats:
    android_lines: int
    backend_lines: int
    total_lines: int
    android_file_count: int
    backend_file_count: int


@dataclass(frozen=True)
class ScreenshotAsset:
    file_name: str
    title: str
    path: Path


def main() -> None:
    repo_root = Path(__file__).resolve().parents[2]
    backend_root = repo_root.parent / "second-bloom-backend"
    output_root = repo_root / "output" / "doc"
    assets_dir = output_root / "assets" / "soft-copyright"
    tmp_dir = repo_root / "tmp" / "soft-copyright"
    output_root.mkdir(parents=True, exist_ok=True)
    assets_dir.mkdir(parents=True, exist_ok=True)
    tmp_dir.mkdir(parents=True, exist_ok=True)

    stats = collect_repo_stats(repo_root=repo_root, backend_root=backend_root)
    assets = ensure_screenshot_assets(assets_dir)

    user_manual_path = output_root / f"{SOFTWARE_NAME}软件著作权用户手册.docx"
    source_program_path = output_root / f"{SOFTWARE_NAME}软件著作权源程序-普通交存.docx"
    info_sheet_path = output_root / "新软件信息采集表-Second Bloom版.docx"

    build_user_manual(
        output_path=user_manual_path,
        repo_root=repo_root,
        backend_root=backend_root,
        stats=stats,
        assets=assets,
    )
    build_source_program_doc(
        output_path=source_program_path,
        repo_root=repo_root,
        backend_root=backend_root,
    )
    convert_docx_to_pdf(user_manual_path)
    convert_docx_to_pdf(source_program_path)
    manual_page_count = pdf_page_count(user_manual_path.with_suffix(".pdf"))
    build_info_sheet(
        output_path=info_sheet_path,
        stats=stats,
        manual_page_count=manual_page_count,
    )
    convert_docx_to_pdf(info_sheet_path)

    print(f"Generated: {user_manual_path}")
    print(f"Generated: {source_program_path}")
    print(f"Generated: {info_sheet_path}")


def collect_repo_stats(repo_root: Path, backend_root: Path) -> RepoStats:
    android_files = sorted((repo_root / "app" / "src" / "main" / "java").rglob("*.kt"))
    backend_files = sorted((backend_root / "app").rglob("*.py"))
    android_lines = sum(count_lines(path) for path in android_files)
    backend_lines = sum(count_lines(path) for path in backend_files)
    return RepoStats(
        android_lines=android_lines,
        backend_lines=backend_lines,
        total_lines=android_lines + backend_lines,
        android_file_count=len(android_files),
        backend_file_count=len(backend_files),
    )


def count_lines(path: Path) -> int:
    with path.open("r", encoding="utf-8") as handle:
        return sum(1 for _ in handle)


def ensure_screenshot_assets(assets_dir: Path) -> list[ScreenshotAsset]:
    assets: list[ScreenshotAsset] = []
    for file_name, title in SCREENSHOT_SPECS:
        path = assets_dir / file_name
        if not path.exists():
            create_placeholder_image(path=path, title=title)
        assets.append(ScreenshotAsset(file_name=file_name, title=title, path=path))
    return assets


def create_placeholder_image(path: Path, title: str) -> None:
    width, height = 1440, 3120
    image = Image.new("RGB", (width, height), color=(247, 243, 232))
    draw = ImageDraw.Draw(image)
    title_font = load_font(84)
    body_font = load_font(44)
    small_font = load_font(34)

    draw.rectangle([(96, 96), (width - 96, height - 96)], outline=(214, 175, 90), width=8)
    draw.rounded_rectangle(
        [(140, 220), (width - 140, 760)],
        radius=42,
        fill=(252, 248, 240),
        outline=(201, 124, 64),
        width=6,
    )
    draw.text((190, 310), SOFTWARE_NAME, fill=(92, 60, 43), font=title_font)
    draw.text((190, 470), title, fill=(141, 79, 49), font=title_font)
    draw.text(
        (190, 650),
        "截图待采集。当前文档已使用占位图，后续可直接覆盖同名文件后重跑生成脚本。",
        fill=(92, 60, 43),
        font=body_font,
    )

    top = 980
    for offset, line in enumerate(
        [
            "建议从 Android 模拟器或真机导出 PNG 截图。",
            "文件名保持不变即可自动替换进用户手册。",
            "推荐顺序：灵感空间 -> 上传识别 -> 方案页 -> 真图编辑 -> 最终效果图。",
        ]
    ):
        draw.rounded_rectangle(
            [(160, top + offset * 260), (width - 160, top + offset * 260 + 190)],
            radius=36,
            fill=(255, 255, 255),
            outline=(214, 175, 90),
            width=4,
        )
        draw.text((220, top + 52 + offset * 260), line, fill=(74, 74, 74), font=body_font)

    draw.text(
        (190, height - 300),
        f"{SOFTWARE_NAME} 软著材料占位截图",
        fill=(128, 128, 128),
        font=small_font,
    )
    image.save(path)


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def set_a4_section(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def base_document() -> Document:
    document = Document()
    set_a4_section(document.sections[0])
    styles = document.styles
    styles["Normal"].font.name = "Songti SC"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Songti SC"
    styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Songti SC"
    styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 2"].font.bold = True
    return document


def add_header(section, text: str) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    run.font.size = Pt(9)


def add_footer_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    run.font.size = Pt(9)


def add_title_block(document: Document, title: str, metadata: list[str]) -> None:
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run(title)
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(124, 70, 35)

    for item in metadata:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(item)
        run.font.name = "Songti SC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
        run.font.size = Pt(10.5)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.name = "Songti SC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
        run.font.size = Pt(10.5)


def add_numbered_steps(document: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.75)
        run = paragraph.add_run(f"{index}. {item}")
        run.font.name = "Songti SC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
        run.font.size = Pt(10.5)


def add_body_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.35
    run = paragraph.add_run(text)
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    run.font.size = Pt(10.5)
    run.bold = bold


def add_info_line(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    label_run = paragraph.add_run(f"{label}：")
    label_run.bold = True
    label_run.font.name = "Songti SC"
    label_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    label_run.font.size = Pt(10.5)
    value_run = paragraph.add_run(value)
    value_run.font.name = "Songti SC"
    value_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    value_run.font.size = Pt(10.5)


def add_screenshot_block(
    document: Document,
    asset: ScreenshotAsset,
    image_intro: str,
    entry_points: list[str],
    interface_notes: list[str],
    usage_steps: list[str],
    data_notes: list[str],
) -> None:
    if len(document.paragraphs) > 18:
        document.add_page_break()
    document.add_heading(asset.title, level=2)
    add_info_line(document, "进入方式", image_intro)
    add_bullets(document, [f"进入方式补充：{item}" for item in entry_points])
    add_bullets(document, [f"界面说明：{item}" for item in interface_notes])
    add_bullets(document, [f"使用说明：{item}" for item in usage_steps])
    add_bullets(document, [f"数据与状态：{item}" for item in data_notes])
    add_body_paragraph(
        document,
        f"{asset.title}在整个业务链路中承担承上启下的作用。页面首先向用户明确当前所处阶段，再围绕该阶段只暴露必要的交互入口，避免用户在一次改造流程中同时处理过多无关操作。"
    )
    add_body_paragraph(
        document,
        "从状态管理角度看，本页面的显示内容并不是静态写死，而是根据 ViewModel 当前状态、历史记录、是否登录、是否正在生成效果图以及是否存在错误提示等条件动态组合出来，因此页面文案、按钮可用性和返回路径都与真实业务状态保持一致。"
    )
    add_body_paragraph(
        document,
        "从数据来源角度看，页面中的文字、图片、卡片和按钮大致来自三类数据：一类是当前内存中的工作流状态，一类是本地历史快照衍生出的展示结果，一类是由后端接口返回的识别、方案或效果图任务数据。三类数据相互配合，保证用户可以从单次改造操作平滑过渡到长期历史沉淀与再次编辑。"
    )
    add_body_paragraph(
        document,
        "从输入角度看，页面接收的主要信息通常包括当前图片、当前分析结果、当前方案、当前登录状态或当前历史快照。不同页面虽然输入内容不同，但都遵循“只展示当前阶段真正需要的信息”的原则，因此用户不必在同一界面处理识别、方案、成图和同步的全部细节。"
    )
    add_body_paragraph(
        document,
        "从输出角度看，页面并不只是展示结果，还会把用户的选择转换成下一步业务动作，例如重新选图、继续低置信度流程、生成方案、创建效果图任务、发布灵感或同步历史。也正因为如此，页面按钮背后都与明确的数据写入、接口调用或状态切换相对应。"
    )
    add_body_paragraph(
        document,
        "从页面关系角度看，本页面与前后页面形成清晰链路：前一页负责提供进入条件，当前页负责完成本阶段处理，后一页负责承接处理结果。这样的结构使系统既适合首次使用者一步步跟随，也适合已有历史用户从任意中间页面恢复操作。"
    )

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(asset.title)
    run.bold = True
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    run.font.size = Pt(10.5)

    document.add_picture(str(asset.path), width=Inches(5.35))
    picture_paragraph = document.paragraphs[-1]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_user_manual(
    output_path: Path,
    repo_root: Path,
    backend_root: Path,
    stats: RepoStats,
    assets: list[ScreenshotAsset],
) -> None:
    document = base_document()
    add_header(document.sections[0], f"{SOFTWARE_NAME} {VERSION} 用户手册")
    add_footer_page_number(document.sections[0])
    add_title_block(
        document=document,
        title=f"{SOFTWARE_NAME}软件著作权用户手册",
        metadata=[
            f"文档名称：{SOFTWARE_NAME}软件著作权用户手册",
            f"文档版本：{VERSION}",
            f"软件版本：{VERSION}",
            "文档用途：用于说明 Second Bloom 旧衣识别、改造方案生成、真图编辑、效果图查看、灵感发布与历史同步等主要能力。",
            f"著作权人：{COPYRIGHT_OWNER}",
        ],
    )

    document.add_heading("1. 软件概述", level=1)
    add_body_paragraph(
        document,
        "Second Bloom旧衣新生系统是一套面向旧衣再设计与数字化改造场景的一体化软件系统。系统由 Android 客户端与 Python 后端服务共同组成，围绕用户从浏览灵感、上传旧衣图片、识别服装属性、确认识别结果、生成改造方案、按选定方案发起真图编辑、查看最终效果图、发布灵感作品到登录后同步历史记录的完整闭环进行设计。"
    )
    add_body_paragraph(
        document,
        "本系统当前已经具备真实的界面流程、状态管理、网络交互、异步效果图任务、本地历史快照及登录后云端历史同步能力。移动端承担用户交互、图片选择、识别确认、方案查看、最终效果图浏览和灵感发布等功能；后端承担服装识别、方案生成、效果图任务创建与轮询、用户资料接口、云端历史读写以及对象存储管理等能力。"
    )
    add_info_line(document, "适用对象", "有旧衣改造需求的普通用户、希望沉淀改造灵感的创作者，以及需要在多设备间同步改造历史的登录用户。")
    add_info_line(document, "部署形态", "Android 客户端 + Python API 服务 + 云端数据库与对象存储。")
    add_info_line(document, "客户端源码规模", f"{stats.android_lines} 行，主要位于 app/src/main/java。")
    add_info_line(document, "后端源码规模", f"{stats.backend_lines} 行，主要位于 second-bloom-backend/app。")
    add_info_line(document, "核心能力", "旧衣图片识别、改造意图选择、改造方案生成、真图编辑、最终效果图查看、灵感发布、本地历史记录、登录与云端同步。")

    document.add_heading("1.1 主要功能", level=2)
    add_bullets(
        document,
        [
            "灵感空间：展示系统内置改造灵感以及用户发布的最新改造成果。",
            "上传识别：从相册选择旧衣图片或加载演示样例，发起识别请求。",
            "识别确认：在低置信度或复杂背景场景下人工核对衣物类型、颜色、材质、风格和瑕疵。",
            "方案生成：根据改造意图和补充偏好生成多套改造方案。",
            "真图编辑：对已选方案追加廓形、长度、领口、袖型与保真策略等微调指令。",
            "最终效果图：轮询异步生成任务并展示改造前、改造后与对比图。",
            "数字衣橱与可持续星球：根据本地保存的分析记录和方案记录衍生历史内容。",
            "我的主页、登录与账号：展示个人资料、语言设置、登录状态与账号中心入口。",
            "云同步：在登录后通过 /me 与 /me/history 系列接口同步历史快照。",
        ],
    )

    document.add_heading("2. 系统组成", level=1)
    add_body_paragraph(
        document,
        "系统当前采用一体化交付思路，前端与后端分别保留独立代码仓，但在业务能力上形成同一个软件系统。Android 端负责所有用户可见页面、状态流转和本地历史；后端负责 AI 识别与生成、对象存储、效果图任务生命周期、用户身份校验以及云端历史接口。"
    )
    add_bullets(
        document,
        [
            "Android 客户端：基于 Kotlin、Jetpack Compose、Navigation Compose 和 ViewModel 构建，提供完整移动端 UI 与本地状态管理。",
            "后端服务：基于 Python 3.12、FastAPI、Pydantic 与 SQLAlchemy 构建，暴露识别、方案、效果图任务、用户资料与历史同步接口。",
            "对象存储：用于保存上传的原始衣物图片和最终效果图资产，文档中对应 beforeImage、afterImage 与 comparisonImage。",
            "数据库：用于持久化分析请求、方案请求、效果图任务、用户资料和云端历史修订信息。",
            "认证与同步：移动端通过 Clerk 配置接入登录，后端通过 JWT 校验保护 /me 系列接口。",
        ],
    )

    document.add_heading("2.1 前后端协同接口说明", level=2)
    add_bullets(
        document,
        [
            "POST /analyze-garment：上传旧衣图片，返回服装分析结果。",
            "POST /generate-remodel-plans：提交确认后的分析结果和改造偏好，返回多套改造方案。",
            "POST /generate-remodel-preview-jobs：基于单个确认方案创建最终效果图任务。",
            "GET /remodel-preview-jobs/{previewJobId}：查询最终效果图任务状态与成图资产。",
            "GET /me：读取登录用户的基础资料。",
            "GET /me/history：读取用户云端历史快照。",
            "POST /me/history/bootstrap：首次把本地历史引导到云端。",
            "PUT /me/history：更新已有云端历史快照。",
        ],
    )
    add_body_paragraph(
        document,
        "上述接口均已在当前代码与文档中落地：Android 端通过 RealRemodelApi 与 HistorySyncHttpClient 访问这些接口，后端通过 app/api/routes 下的 analyze、plans、previews 与 me 路由对外暴露能力。为了保持文档稳定，本手册只描述接口作用与业务意义，不直接展示 OpenAPI 或 JSON 详情。"
    )

    document.add_heading("3. 客户端功能说明", level=1)
    screenshot_map = {asset.file_name: asset for asset in assets}

    add_screenshot_block(
        document,
        screenshot_map["01_inspiration.png"],
        image_intro="应用启动后默认进入灵感空间首页，底部导航当前高亮“灵感空间”。",
        entry_points=[
            "应用启动自动进入 inspiration 路由。",
            "在底部导航点击“灵感空间”也可返回该页面。",
        ],
        interface_notes=[
            "页面顶部展示 Second Bloom 品牌标题与灵感说明文案。",
            "中部以卡片形式展示系统内置改造灵感以及用户发布的作品。",
            "每张卡片可进入灵感详情页，查看改造前后图片、点赞、收藏与评论入口。",
        ],
        usage_steps=[
            "用户可先浏览系统推荐的旧衣改造灵感。",
            "点击“开始 AI 改制”可直接跳转到上传识别流程。",
            "点击某条灵感卡片可进入详情页查看完整内容。",
        ],
        data_notes=[
            "灵感空间会结合本地发布记录 prepend 最新用户作品。",
            "点赞、收藏和评论交互会写入本地 inspiration engagement 记录。",
        ],
    )

    add_screenshot_block(
        document,
        screenshot_map["02_upload_demo.png"],
        image_intro="从灵感空间点击“开始 AI 改制”，或点击底部中间加号按钮后进入上传识别页面。",
        entry_points=[
            "通过 camera_recognition 路由进入上传识别页面。",
            "可选择真实相册图片，也可加载 DemoScenario 演示样例。",
        ],
        interface_notes=[
            "页面显示上传、确认、方案三段式工作流条。",
            "主上传卡片包含图片选择区和开始识别按钮。",
            "底部展示三种演示场景：正常识别、复杂背景、异常提示。",
        ],
        usage_steps=[
            "用户点击“选择照片”从相册选取一张旧衣图片。",
            "也可以直接点击某个示例场景，快速切换到演示图像。",
            "确认图片后点击“开始识别”，系统进入识别阶段。",
        ],
        data_notes=[
            "SelectedImage 记录图片 URI、文件名、MIME 类型和大小。",
            "演示场景使用 DemoScenario 转换为 demo:// URI 以支持演示流程。",
        ],
    )

    add_screenshot_block(
        document,
        screenshot_map["03_low_confidence.png"],
        image_intro="当识别置信度偏低或背景复杂时，系统停留在低置信度提示状态。",
        entry_points=[
            "复杂背景样例可稳定触发该界面。",
            "真实图片识别后若 backgroundComplexity 为 HIGH 或 confidence 低于阈值，也会出现提示。",
        ],
        interface_notes=[
            "页面会展示提醒卡片，说明当前识别存在不确定性。",
            "用户可以重新选择图片，也可以继续使用当前识别结果。",
            "工作流仍停留在“确认”阶段，等待人工判断。",
        ],
        usage_steps=[
            "用户阅读提示信息并核对当前衣物是否可继续使用。",
            "若识别偏差较大，可重新选择图片重新识别。",
            "若整体结果可接受，可点击继续，进入识别摘要确认流程。",
        ],
        data_notes=[
            "该状态由 RemodelViewModel 根据置信度与背景复杂度自动计算。",
            "继续操作不会丢失当前分析结果，而是把 stage 切换为 AnalysisReady。",
        ],
    )

    add_screenshot_block(
        document,
        screenshot_map["04_analysis_summary.png"],
        image_intro="识别成功后，系统在方案页展示识别摘要，并允许用户人工修改关键字段。",
        entry_points=[
            "上传识别成功后点击“进入方案页”。",
            "或在低置信度提示后选择继续，进入方案页。",
        ],
        interface_notes=[
            "页面上方展示摘要、目标选择和方案三段式工作流。",
            "识别摘要区域显示衣物类型、颜色、材质、风格、瑕疵等内容。",
            "可展开编辑器对识别结果做人工修正。",
        ],
        usage_steps=[
            "用户先阅读识别摘要，判断结果是否准确。",
            "如需修正，可展开编辑区域修改文字内容。",
            "确认无误后选择改造目标并填写补充偏好。",
        ],
        data_notes=[
            "GarmentAnalysis 对象会在 analysis 与 draftAnalysis 两份状态中保存。",
            "编辑操作只修改 draftAnalysis，确保原始识别结果仍可追溯。",
        ],
    )

    add_screenshot_block(
        document,
        screenshot_map["05_plan_generation.png"],
        image_intro="在方案页选择改造目标并点击生成后，系统返回多套改造方案。",
        entry_points=[
            "用户在方案页选择改造意图，例如日常改造、特殊场合、创意 DIY 或尺码调整。",
            "填写偏好后点击“生成改制方案”。",
        ],
        interface_notes=[
            "每张方案卡片展示标题、摘要、难度、耗时、材料和关键步骤。",
            "已生成方案会自动保存到历史记录，供数字衣橱和个人主页复用。",
            "方案卡片同时提供进入真图编辑和查看效果图状态的入口。",
        ],
        usage_steps=[
            "用户比较多套方案后，挑选最符合自己需求的一套。",
            "点击进入真图编辑，继续微调该方案。",
            "若已有成图记录，也可直接查看最终效果图状态。",
        ],
        data_notes=[
            "RemodelPlan 统一描述方案标题、摘要、难度、材料、预计耗时和步骤。",
            "生成成功后会写入 SavedPlanGenerationRecord 并刷新本地历史。",
        ],
    )

    add_screenshot_block(
        document,
        screenshot_map["06_preview_editor.png"],
        image_intro="选中方案后进入真图编辑页，对已生成方案追加细化指令。",
        entry_points=[
            "点击方案卡片中的“进入真图编辑”按钮。",
            "从最终效果图页点击“继续编辑该方案”也可回到本页。",
        ],
        interface_notes=[
            "页面展示当前原图和选中的方案摘要。",
            "用户可调节整体廓形、衣长、领口、袖型和保真策略。",
            "还可以填写额外微调说明，例如保留旧衣质感或不要改变背景。",
        ],
        usage_steps=[
            "用户先确认当前原始衣物图片和目标方案无误。",
            "根据需要选择微调选项，补充一到两句具体说明。",
            "点击“生成最终效果图”，系统立即创建异步成图任务。",
        ],
        data_notes=[
            "PreviewEditOptions 负责承载各项微调字段。",
            "客户端通过 createPreviewJob 接口把 planId、analysisId 和 tuning 信息发送给后端。",
        ],
    )

    add_screenshot_block(
        document,
        screenshot_map["07_preview_result.png"],
        image_intro="生成最终效果图后，系统在专用结果页展示任务状态和成图结果。",
        entry_points=[
            "从真图编辑页提交生成后自动进入结果页。",
            "也可从方案卡片点击“查看最终效果图”进入。",
        ],
        interface_notes=[
            "结果页支持展示排队中、生成中、失败、已完成等不同状态。",
            "完成后可集中查看改造前图片、改造后图片和对比图。",
            "成功结果可发布到灵感空间，也可继续编辑或返回方案页。",
        ],
        usage_steps=[
            "当任务仍在运行时，用户可停留本页等待轮询刷新。",
            "成图完成后，用户比较前后效果并决定是否发布。",
            "若结果过滤或失败，可返回编辑页再次调整并重新发起生成。",
        ],
        data_notes=[
            "PreviewJob 包含任务级状态，PlanPreviewResult 记录单方案渲染结果。",
            "轮询接口返回 beforeImage、afterImage 与 comparisonImage 等资产信息。",
        ],
    )

    add_screenshot_block(
        document,
        screenshot_map["08_wardrobe.png"],
        image_intro="底部导航切换到“数字衣橱”后，可查看已识别和已生成方案的历史条目。",
        entry_points=[
            "点击底部导航中的“数字衣橱”。",
            "当历史中有保存记录时会自动生成衣橱卡片。",
        ],
        interface_notes=[
            "页面按衍生类别和卡片方式展示历史改造对象。",
            "用户可从衣橱条目重新打开曾经保存的方案结果。",
            "页面保留回到 AI 改制流程的快速入口。",
        ],
        usage_steps=[
            "用户浏览历史衣物与已保存方案。",
            "点击条目可恢复某次方案生成上下文，继续编辑。",
            "也可再次进入 AI 改制流程发起新的识别任务。",
        ],
        data_notes=[
            "数字衣橱不是独立数据库，而是对本地历史快照的衍生展示。",
            "其内容由 SavedAnalysisRecord 与 SavedPlanGenerationRecord 计算得到。",
        ],
    )

    add_screenshot_block(
        document,
        screenshot_map["09_planet.png"],
        image_intro="底部导航切换到“可持续星球”后，可查看由历史记录推导出的环保影响信息。",
        entry_points=[
            "点击底部导航中的“可持续星球”。",
            "该页面依赖已有识别与方案历史记录自动生成内容。",
        ],
        interface_notes=[
            "页面集中展示旧衣再利用与改造累计带来的环保效果。",
            "用户可从本页再次进入改造流程，继续扩充历史记录。",
            "整体信息由轻量衍生模型计算，不需要单独维护表单。",
        ],
        usage_steps=[
            "用户查看当前累计的改造次数与环保相关摘要。",
            "理解历史记录如何转化为可持续展示内容。",
            "如需增加记录，可回到上传识别入口创建新的改造流程。",
        ],
        data_notes=[
            "页面数据来自 deriveSustainabilityImpactSummary 等衍生方法。",
            "系统保持轻量化设计，不引入额外数据库表或重量级分析框架。",
        ],
    )

    add_screenshot_block(
        document,
        screenshot_map["10_profile.png"],
        image_intro="底部导航切换到“我的主页”后，可查看个人资料、作品、语言设置与登录状态。",
        entry_points=[
            "点击底部导航中的“我的主页”。",
            "未登录与已登录状态在同一页面上采用不同展示文案。",
        ],
        interface_notes=[
            "页面顶部显示头像、昵称和当前账号提示语。",
            "中部提供作品、已保存内容、语言选择和账号卡片。",
            "登录后可进入账号中心，未登录时则展示登录入口。",
        ],
        usage_steps=[
            "用户可在本页切换中英文界面。",
            "查看最近作品与已保存内容，理解历史记录的聚合展示。",
            "根据当前状态点击登录、账号中心或退出登录等操作。",
        ],
        data_notes=[
            "本地历史会驱动个人主页中的作品墙、最近动态和数据摘要。",
            "语言设置通过 AppPreferencesRepository 持久化保存。",
        ],
    )

    add_screenshot_block(
        document,
        screenshot_map["11_auth.png"],
        image_intro="从个人主页点击登录入口后进入登录页面，登录成功后可进入账号中心并启动云同步。",
        entry_points=[
            "个人主页点击“登录 / 注册”进入 auth 路由。",
            "已登录用户点击“账号中心”进入 account 路由。",
        ],
        interface_notes=[
            "登录页承接 Clerk 提供的认证界面。",
            "账号页展示当前用户身份和相关说明。",
            "认证完成后系统会自动拉起本地历史与云端历史的引导和推送流程。",
        ],
        usage_steps=[
            "用户输入账号信息或使用已配置方式登录。",
            "登录成功后返回个人主页或进入账号中心。",
            "随后系统可通过 /me 与 /me/history 接口同步历史快照。",
        ],
        data_notes=[
            "移动端通过 SecondBloomClerkConfig 读取发布密钥配置。",
            "未登录状态下历史仍保存在本地，登录后才会触发云端同步。",
        ],
    )

    document.add_heading("4. 服务端支撑能力", level=1)
    add_body_paragraph(
        document,
        "后端系统位于 sibling 仓库 second-bloom-backend 中，当前已具备可运行的 FastAPI 应用入口与多组 API 路由。系统通过 app/main.py 创建应用实例，并挂载 health、analyze、plans、previews 和 me 路由。业务层主要集中在 app/api/services.py、app/services/*、app/providers/* 和 app/db/*。"
    )
    add_bullets(
        document,
        [
            "服装识别服务：接收上传图片，生成结构化 GarmentAnalysis 结果。",
            "方案生成服务：根据确认后的识别信息和用户偏好返回多套 RemodelPlan。",
            "异步效果图任务：创建 preview job，保存状态并返回任务查询接口。",
            "最终成图资产：管理改造前、改造后与对比图的存储和 URL 下发。",
            "用户资料与云同步：在登录状态下提供 /me、/me/history、/me/history/bootstrap 和 /me/history 更新接口。",
            "认证校验：后端通过 Clerk JWT 验证保护用户相关路由。",
        ],
    )
    add_body_paragraph(
        document,
        "在生产部署中，后端 README 说明系统使用 Vercel 作为部署平台，数据库可由 Neon 提供，对象存储由 Vercel Blob 提供。对于本手册而言，这些技术组件的意义在于确保旧衣图片、最终效果图和用户历史等核心业务数据能够在客户端与服务端之间稳定传递、持久保存并按需读取。"
    )

    document.add_heading("4.1 核心业务对象说明", level=2)
    add_bullets(
        document,
        [
            "服装分析结果：用于表示识别后的衣物类型、颜色、材质、风格、背景复杂度、置信度和瑕疵列表。",
            "改造意图：用于描述本次改造的目标方向，例如日常穿着、特殊场合、创意 DIY 或尺码调整。",
            "改造方案：用于表示某次生成得到的一套完整改造建议，包括摘要、难度、材料、预计耗时与步骤。",
            "效果图任务：用于表示一次最终效果图生成的任务状态、结果数量与资产集合。",
            "本地或云端历史快照：用于汇总分析记录、方案记录、已发布作品和灵感互动记录。",
        ],
    )
    add_body_paragraph(
        document,
        "这些业务对象在前后端之间并不是孤立存在的。客户端以中文化的界面形式展示这些对象，后端以结构化模型和接口响应形式返回这些对象，最终共同构成“识别 - 方案 - 成图 - 发布 - 同步”的主链路。对于软著材料编写来说，用统一术语描述这些对象，有助于保持用户手册、信息采集表和源码交存之间的一致性。"
    )

    document.add_heading("4.2 后端任务与存储支撑说明", level=2)
    add_bullets(
        document,
        [
            "上传识别阶段：后端接收旧衣图片并生成结构化分析结果。",
            "方案阶段：后端根据确认后的分析结果和用户偏好生成多套改造方案。",
            "效果图阶段：后端创建预览任务并通过轮询接口回传任务状态。",
            "资产阶段：后端维护改造前、改造后和对比图等图像资产地址。",
            "用户阶段：后端校验登录用户身份并提供云端历史接口。",
        ],
    )
    add_body_paragraph(
        document,
        "后端 README 已明确说明生产环境中的数据库和对象存储职责。数据库承担分析请求、方案请求、任务状态和用户历史等结构化数据的保存；对象存储承担上传图片和效果图资产的保存。这样做的直接效果是：即使最终效果图任务需要较长时间处理，客户端也可以通过任务轮询和资产地址稳定获取处理结果。"
    )
    add_body_paragraph(
        document,
        "与此同时，后端并未把本地历史的全部逻辑迁移到云端，而是保留“本地优先、登录后同步”的边界。这样既能确保游客模式下系统照常可用，也让登录用户获得跨设备恢复能力，符合当前代码中的轻量化实现边界。"
    )

    document.add_heading("5. 典型使用流程", level=1)
    process_sections = [
        (
            "5.1 游客改造流程",
            [
                "用户打开应用后先进入灵感空间，浏览改造案例。",
                "点击开始 AI 改制，进入上传识别页面。",
                "从相册选择一张旧衣图片并发起识别。",
                "识别成功后进入方案页，核对识别摘要。",
                "选择改造意图、填写偏好并生成多套方案。",
                "挑选一套方案进入真图编辑，调整微调选项。",
                "提交后进入最终效果图页，等待任务完成并查看结果。",
            ],
        ),
        (
            "5.2 低置信度修正流程",
            [
                "用户选择复杂背景或识别难度较高的图片。",
                "系统返回低置信度提示，提醒结果可能存在偏差。",
                "用户可重新选图，也可继续使用当前结果。",
                "进入方案页后，用户对衣物类型、颜色、材质、风格和瑕疵进行人工修改。",
                "修正完成后再次生成方案，确保后续改造结果更贴合真实衣物。",
            ],
        ),
        (
            "5.3 最终效果图生成流程",
            [
                "用户在方案列表中挑选目标方案并进入真图编辑页。",
                "根据想保留的旧衣特征设置廓形、长度、领口、袖型和保真策略。",
                "系统把 analysisId、planId 和微调参数发送给后端。",
                "后端创建异步效果图任务，客户端开始轮询状态。",
                "任务完成后结果页展示改造前、改造后和对比图。",
                "若任务过滤或失败，用户可以回到真图编辑页再次调整。",
            ],
        ),
        (
            "5.4 发布到灵感空间流程",
            [
                "用户在最终效果图页确认成图结果可接受。",
                "点击发布按钮后，系统把当前结果写入已发布改造记录。",
                "返回灵感空间后，用户发布的最新作品会优先展示在顶部。",
                "其他灵感详情页中也可以继续触发点赞、收藏和评论等互动。",
            ],
        ),
        (
            "5.5 登录后云同步流程",
            [
                "用户在我的主页中点击登录入口，完成认证。",
                "登录成功后，系统读取用户资料并初始化云端历史接口。",
                "首次同步时执行 bootstrap，把本地历史推送为云端初始快照。",
                "后续识别、方案生成、发布等操作产生的新历史会被后台推送到云端。",
                "当用户在其他设备登录同一账号时，可拉取云端快照恢复历史。",
            ],
        ),
        (
            "5.6 历史恢复与再次编辑流程",
            [
                "用户打开数字衣橱或我的主页，浏览过去保存的识别记录和方案记录。",
                "从历史条目中选择某次方案生成记录，恢复到当前工作台上下文。",
                "系统把对应的衣物分析结果、改造意图和方案列表重新装载到当前状态。",
                "用户可基于历史内容重新进入真图编辑页或重新查看最终效果图结果。",
                "这一流程保证系统不只是一次性生成工具，还能支持历史内容回看与再创作。",
            ],
        ),
    ]
    for title, steps in process_sections:
        document.add_heading(title, level=2)
        add_numbered_steps(document, steps)

    document.add_heading("6. 异常与提示说明", level=1)
    exception_items = [
        "图片无效：当用户未选择图片、上传文件类型不正确或图片无法解析时，客户端会提示先选择可识别的旧衣图片。",
        "网络失败：当识别、方案生成或历史同步请求失败时，界面会展示网络请求失败提示，用户可重新发起操作。",
        "模型结果异常：当服务端返回空结果或不可识别响应时，客户端会给出“服务端返回了无法识别的结果”等错误说明。",
        "预览超时或过滤：最终效果图生成较慢、结果被视觉质量校验过滤或资产缺失时，结果页与编辑页都会展示当前状态说明。",
        "未登录同步限制：在未登录状态下，本地历史仍可保存和使用，但不会调用 /me 与 /me/history 系列云同步接口。",
    ]
    for item in exception_items:
        add_body_paragraph(document, item)

    document.add_heading("6.1 关键状态说明", level=2)
    add_bullets(
        document,
        [
            "背景复杂度状态：系统当前使用 low 与 high 两类背景复杂度区分普通背景与复杂背景。",
            "识别置信度状态：当识别置信度低于阈值时，页面会进入低置信度确认阶段。",
            "效果图任务状态：任务级状态包括 queued、running、completed、completed_with_failures、failed 与 expired。",
            "效果图渲染状态：单方案结果状态包括 queued、running、completed、failed 与 filtered。",
            "同步状态：未登录时只保留本地历史，登录后再执行 bootstrap 与后续 push 或 refresh。",
        ],
    )
    add_body_paragraph(
        document,
        "这些状态贯穿上传识别、方案生成、真图编辑、结果轮询和云端同步等多个流程。客户端通过 ViewModel 管理状态流，后端通过接口和任务对象返回对应状态，因此用户在不同页面看到的提示、进度条和可用操作都与这些状态直接相关。"
    )

    document.add_heading("6.2 本地历史与云端历史说明", level=2)
    add_body_paragraph(
        document,
        "系统采用轻量历史快照策略组织历史数据。客户端通过 remodel_history.json 保存本地快照，快照至少包含分析记录、方案记录、已发布改造记录和灵感互动记录。数字衣橱、可持续星球和我的主页都依赖这些历史数据实时衍生展示，因此一次识别、一次方案生成或一次发布完成后，系统都会刷新相关页面。"
    )
    add_bullets(
        document,
        [
            "分析记录：保存旧衣原图和结构化识别结果，用于后续方案生成与历史回看。",
            "方案记录：保存改造意图、用户偏好和生成的多套方案，用于恢复工作台上下文。",
            "已发布作品：保存可展示的最终效果图结果，用于灵感空间和个人主页展示。",
            "灵感互动：保存点赞、收藏和评论等互动行为。",
            "云端快照：在登录状态下通过 schemaVersion、revision 与 snapshot 结构同步到后端。",
        ],
    )

    document.add_heading("6.3 识别、方案与效果图字段说明", level=2)
    add_bullets(
        document,
        [
            "识别结果字段重点包括 analysisId、衣物类型、颜色、材质、风格、瑕疵列表、背景复杂度、置信度和提示信息。",
            "方案结果字段重点包括 planId、方案标题、方案摘要、难度、材料、预计耗时和步骤列表。",
            "真图编辑字段重点包括廓形、衣长、领口、袖型、保真策略和补充说明。",
            "效果图任务字段重点包括 previewJobId、任务状态、请求方案数量、结果列表以及图像资产地址。",
            "历史快照字段重点包括 analyses、planGenerations、publishedRemodels 和 inspirationEngagements。",
        ],
    )
    add_body_paragraph(
        document,
        "虽然用户在界面上看到的是中文化文案和卡片，但这些字段在内部都对应稳定的数据模型。客户端据此恢复页面状态，后端据此维持接口契约，因此这些字段说明既能帮助理解软件内部结构，也能解释为什么不同页面之间的数据能够连贯流转。"
    )

    document.add_heading("7. 运行环境", level=1)
    environment_rows = [
        ("开发硬件环境", "开发电脑：8GB 以上内存、可用磁盘空间 5GB 以上。"),
        ("客户端运行环境", "Android 智能手机，支持图片选择、网络访问和本地文件存储。"),
        ("客户端操作系统", CLIENT_MIN_OS),
        ("服务端运行环境", BACKEND_RUNTIME),
        ("主要开发工具", "Android Studio、Gradle、JDK 17、Jetpack Compose、Python 3.12、FastAPI、SQLAlchemy、Vercel。"),
        ("编程语言", "Kotlin、Python。"),
        ("客户端支撑环境", "Android 系统环境、网络环境、本地文件存储环境。"),
        ("服务端支撑环境", "HTTP 服务环境、对象存储、数据库与用户认证校验环境。"),
    ]
    for label, value in environment_rows:
        add_info_line(document, label, value)

    document.add_heading("7.1 客户端主要页面与路由对照", level=2)
    add_bullets(
        document,
        [
            "inspiration：灵感空间首页。",
            "inspiration_detail/{itemId}：灵感详情页。",
            "camera_recognition：上传识别页。",
            "plan：方案页。",
            "preview_editor/{planId}：真图编辑页。",
            "preview_result/{planId}：最终效果图页。",
            "wardrobe：数字衣橱页。",
            "planet：可持续星球页。",
            "profile：我的主页。",
            "auth：登录页。",
            "account：账号中心页。",
        ],
    )
    add_body_paragraph(
        document,
        "这些路由均已在客户端 Screen 定义中固化，并由 MainScreen 中的 NavHost 统一管理。手册中的截图和步骤围绕这些路由组织，因此页面命名、跳转路径和界面职责都可以与当前代码直接对应。"
    )

    document.add_heading("7.2 后端主要模块说明", level=2)
    add_bullets(
        document,
        [
            "app/api/routes：对外暴露 health、analyze、plans、previews 与 me 等接口路由。",
            "app/api/services.py：聚合识别、方案生成、效果图任务、用户资料和存储访问等核心编排逻辑。",
            "app/services：按识别、方案、效果图任务与渲染等职责拆分服务模块。",
            "app/providers：封装 Qwen 识别、方案、图像编辑、视觉质量校验等模型能力。",
            "app/db：封装对象存储、数据库模型、仓储与配置逻辑。",
        ],
    )
    add_body_paragraph(
        document,
        "这一拆分方式保证了前后端一体系统在功能上保持统一，在实现上又能把界面交互、HTTP 接口、模型能力接入、数据持久化和用户历史同步等问题清晰分层。对于软件著作权文档而言，这些模块说明能够帮助审查者理解该系统并非静态演示界面，而是由多个协同模块共同组成的完整软件系统。"
    )

    document.add_heading("7.3 页面与数据来源对照", level=2)
    add_bullets(
        document,
        [
            "灵感空间与灵感详情页：来源于静态展示模型与已发布改造记录、互动记录的组合。",
            "上传识别页：来源于当前所选图片和演示场景数据。",
            "方案页：来源于当前草稿识别结果、改造意图和方案生成结果。",
            "真图编辑页：来源于当前选中方案、原始图片和微调选项。",
            "最终效果图页：来源于效果图任务状态、结果对象和图像资产地址。",
            "数字衣橱、可持续星球、我的主页：来源于本地历史快照与其衍生模型。",
            "登录与账号页：来源于 Clerk 登录状态和后端用户资料接口。",
        ],
    )
    add_body_paragraph(
        document,
        "这种“页面围绕数据模型组织”的方式让系统既适合日常使用，也适合作为软件著作权申请材料说明对象。审查者可以通过页面理解软件功能，再通过字段和模块理解这些功能并非孤立页面，而是有明确状态来源和数据支撑的完整系统。"
    )

    document.add_heading("7.4 部署与运行支撑说明", level=2)
    add_body_paragraph(
        document,
        "从运行角度看，客户端安装在 Android 设备上，负责所有用户交互；后端部署在云端服务环境中，负责识别、方案和效果图任务；数据库和对象存储则为服务端提供结构化数据与图像资产存储支持。用户登录能力通过外部认证服务接入，但具体登录后的页面逻辑、资料读取和历史同步仍由本系统自行控制。"
    )
    add_body_paragraph(
        document,
        "这种部署方式使系统既支持纯本地游客模式，也支持联网增强模式。游客模式下用户仍可完成旧衣识别、方案生成和本地历史沉淀；联网增强模式下用户可获得真实服务端识别结果、异步成图结果和云端历史同步能力。两种模式在当前代码中都已有对应实现，因此能够作为同一软件系统中的不同运行形态进行说明。"
    )

    document.add_heading("8. 结论", level=1)
    add_body_paragraph(
        document,
        "综上，Second Bloom旧衣新生系统已经形成由 Android 客户端与 Python 后端共同支撑的一体化软件系统。系统能够围绕旧衣识别、改造方案生成、真图编辑、最终效果图查看、灵感发布和历史同步等核心业务形成完整闭环，并已经在当前代码仓与后端仓中实现对应能力。"
    )
    add_body_paragraph(
        document,
        "本手册仅描述截至当前版本已经实现并可在代码和接口文档中核实的功能，不包含未来规划功能，也不引入与实际代码不一致的业务叙述，可作为本系统申请软件著作权时的文档鉴别材料使用。"
    )

    document.save(output_path)


def build_source_program_doc(output_path: Path, repo_root: Path, backend_root: Path) -> None:
    ordered_files: list[tuple[str, Path]] = []
    for relative_path in ANDROID_FILE_ORDER:
        ordered_files.append((relative_path, repo_root / relative_path))
    for relative_path in BACKEND_FILE_ORDER:
        ordered_files.append((relative_path, backend_root / relative_path))

    source_lines: list[str] = []
    line_number = 1
    for relative_path, file_path in ordered_files:
        if not file_path.exists():
            continue
        source_lines.append(format_program_line(line_number, f"// File: {relative_path}"))
        line_number += 1
        with file_path.open("r", encoding="utf-8") as handle:
            for raw_line in handle:
                source_lines.append(format_program_line(line_number, raw_line.rstrip("\n")))
                line_number += 1

    if len(source_lines) < 3000:
        raise RuntimeError("Source line count is too small to produce the required 60 pages.")

    first_chunk = source_lines[:1500]
    last_chunk = source_lines[-1500:]
    selected_lines = first_chunk + last_chunk

    document = base_document()
    set_a4_section(document.sections[0])
    add_header(document.sections[0], f"{SOFTWARE_NAME} {VERSION} 源程序交存")
    add_footer_page_number(document.sections[0])

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{SOFTWARE_NAME}软件著作权源程序-普通交存")
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    run.font.size = Pt(16)
    run.bold = True

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run = meta.add_run(
        f"软件版本：{VERSION}    交存方式：前 30 页 + 后 30 页    每页 50 行连续页码"
    )
    meta_run.font.name = "Songti SC"
    meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    meta_run.font.size = Pt(10.5)

    page_count = math.ceil(len(selected_lines) / 50)
    for page_index in range(page_count):
        chunk = selected_lines[page_index * 50:(page_index + 1) * 50]
        if len(chunk) < 50:
            chunk += [format_program_line(99999, "")] * (50 - len(chunk))

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run("\n".join(chunk))
        run.font.name = "Menlo"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
        run.font.size = Pt(8.3)

        if page_index != page_count - 1:
            paragraph.runs[-1].add_break(WD_BREAK.PAGE)

    document.save(output_path)


def format_program_line(line_number: int, content: str) -> str:
    if line_number == 99999:
        return " " * 6
    sanitized = content.expandtabs(4)
    return f"{line_number:05d}  {sanitized}"


def build_info_sheet(output_path: Path, stats: RepoStats, manual_page_count: int) -> None:
    repo_root = output_path.parents[2]
    template_path = (
        repo_root.parent
        / "2026"
        / "NyxGuard"
        / "output"
        / "doc"
        / "新软件信息采集表-完成版.docx"
    )
    if not template_path.exists():
        raise FileNotFoundError(f"Info sheet template not found: {template_path}")

    shutil.copyfile(template_path, output_path)
    document = Document(output_path)
    table = document.tables[0]

    main_features = (
        "Second Bloom旧衣新生系统是一套面向旧衣识别、改造方案生成、真图编辑、"
        "最终效果图查看与历史同步的一体化软件系统。系统由 Android 客户端与 Python 后端服务组成，"
        "以用户从浏览灵感、上传旧衣图片、识别服装属性、确认识别结果、选择改造意图、生成多套改造方案、"
        "按选定方案追加微调指令、异步生成最终效果图到发布灵感和同步历史记录的完整流程为主线。"
        "客户端提供灵感空间、上传识别、方案页、真图编辑、最终效果图、数字衣橱、可持续星球、我的主页、"
        "登录与账号中心等界面能力；后端提供服装识别、方案生成、效果图任务创建与查询、用户资料、云端历史读写等接口能力。"
        "系统采用 Kotlin 与 Python 开发，前端基于 Jetpack Compose 与 ViewModel 管理状态，后端基于 FastAPI、Pydantic 和 SQLAlchemy 提供服务。"
        "通过本地历史快照与登录后的云端历史快照机制，系统能够把识别记录、方案记录、已发布作品和灵感互动统一沉淀，"
        "并支撑数字衣橱、可持续展示和跨设备恢复等实际业务场景。"
    )

    technical_features = (
        "1、采用 Android Studio、Gradle、JDK 17、Jetpack Compose 构建移动端界面与导航流程；\n"
        "2、采用 Python 3.12、FastAPI、Pydantic、SQLAlchemy 构建后端识别、方案与效果图任务服务；\n"
        "3、通过本地 JSON 快照实现轻量历史存储，并通过 /me/history 系列接口实现登录后的云端同步；\n"
        "4、通过异步效果图任务机制管理改造前、改造后与对比图等资产，支持状态轮询与失败提示。"
    )

    row_values = {
        1: SOFTWARE_NAME,
        2: "Second Bloom",
        3: VERSION,
        4: "应用软件",
        5: COMPANY_FOUNDING_DATE,
        6: SOFTWARE_COMPLETION_DATE,
        7: SOFTWARE_PUBLICATION_DATE,
        8: FIRST_PUBLICATION_LOCATION,
        9: "开发电脑：CPU 2.0GHz以上；内存8GB以上；可用磁盘空间5GB以上。",
        10: "Android智能手机：内存2GB以上；可用存储空间100MB以上；支持图片选择、网络访问与本地文件存储。",
        11: "macOS",
        12: "Android Studio、Gradle、JDK 17、Jetpack Compose、Python 3.12、FastAPI、SQLAlchemy、Vercel",
        13: "Android 7.0及以上；Linux云端服务环境",
        14: "Android系统环境、HTTP服务环境、数据库、对象存储、用户认证与网络环境",
        15: "Kotlin、Python",
        16: f"{stats.total_lines}行",
        17: "面向旧衣再设计与数字化改造场景，提供识别、方案生成、真图编辑、成图展示与历史沉淀能力。",
        18: "移动互联网、服装数字化改造、AI图像应用、可持续生活方式",
        19: main_features,
        20: technical_features,
        21: f"{manual_page_count}页",
        22: "60页",
    }
    for row_index, value in row_values.items():
        set_cell_text_preserving_style(table.cell(row_index, 1), value)
        set_cell_text_preserving_style(table.cell(row_index, 2), value)

    set_cell_text_preserving_style(table.cell(25, 0), COPYRIGHT_OWNER)
    set_cell_text_preserving_style(table.cell(25, 1), COPYRIGHT_OWNER)
    set_cell_text_preserving_style(table.cell(25, 2), OWNER_CERTIFICATE)
    document.save(output_path)


def set_cell_text_preserving_style(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for extra_paragraph in cell.paragraphs[1:]:
        for run in extra_paragraph.runs:
            run.text = ""


def convert_docx_to_pdf(docx_path: Path) -> None:
    output_dir = docx_path.parent
    profile_dir = Path("/tmp/libreoffice-soft-copyright-profile")
    profile_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            str(resolve_soffice()),
            f"-env:UserInstallation=file://{profile_dir}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )


def resolve_soffice() -> Path:
    candidates = [
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
        Path("/opt/homebrew/Caskroom/libreoffice/26.2.2/LibreOffice.app/Contents/MacOS/soffice"),
        Path("/opt/homebrew/bin/soffice"),
        Path("/usr/local/bin/soffice"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    resolved = shutil.which("soffice")
    if resolved:
        return Path(resolved)
    raise FileNotFoundError("Unable to find LibreOffice soffice executable")


def pdf_page_count(pdf_path: Path) -> int:
    result = subprocess.run(
        ["pdfinfo", str(pdf_path)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    for line in result.stdout.splitlines():
        if line.startswith("Pages:"):
            return int(line.split(":", 1)[1].strip())
    raise RuntimeError(f"Unable to determine page count for {pdf_path}")


if __name__ == "__main__":
    main()
